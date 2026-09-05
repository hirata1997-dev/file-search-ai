import logging
from pathlib import Path
from datetime import datetime
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook
from database import (
    initialize_db,
    save_files_to_db,
    delete_missing_files_from_db,
    count_files_in_db,
    search_by_extension_db,
    search_by_name_db,
    search_by_content_db
)

# スキャン処理や読み込みエラーをログファイルへ記録
logging.basicConfig(
    filename="scanner.log",
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)


def scan_files(target_dir):
    """指定フォルダ以下を再帰的に走査し、ファイル情報と本文を取得する。"""

    files = []

    # 指定フォルダ以下のファイル・フォルダを再帰的に確認
    for path in target_dir.rglob("*"):
        try:
            if path.is_file():

                content = None

                extension = path.suffix.lower()

                # ファイル形式ごとに検索対象となる本文を抽出
                if extension == ".txt":
                    content = read_text_file(path)

                elif extension == ".pdf":
                    content = read_pdf_file(path)

                elif extension == ".docx": 
                    content = read_docx_file(path)

                elif extension == ".xlsx":
                    content = read_excel_file(path)

                # DBへ保存するファイル情報を作成
                file_info = {
                    "name": path.name,
                    "extension": extension,
                    "path": str(path),
                    "size": path.stat().st_size,
                    "modified": datetime.fromtimestamp(path.stat().st_mtime),
                    "content": content
                }

                files.append(file_info)

        except (PermissionError, OSError) as error:
            logging.error("読み込み失敗: %s / 理由: %s", path, error)

    return files


def read_text_file(path):
    """TXTファイルをUTF-8またはCP932で読み込み、本文を返す。"""

    encodings = ["utf-8", "cp932"]

    # UTF-8で失敗した場合はCP932で再試行
    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as file:
                return file.read()

        except UnicodeDecodeError:
            continue
            
    logging.warning(
        "対応している文字コードで読み込めませんでした: %s",
        path
    )

    return None


def read_pdf_file(path):
    """PDFファイルから検索対象となるテキストを抽出する。"""

    try:
        reader = PdfReader(path)

        content = ""

        # PDFをページ単位で読み込み、本文を結合
        for page in reader.pages:
            text = page.extract_text()

            if text:
                content += text + "\n"

        return content

    except Exception as error:
        logging.error(
            "PDF読み込み失敗: %s / 理由: %s",
            path,
            error
        )

        return None


def read_docx_file(path):
    """Wordファイルの段落と表から検索対象となるテキストを抽出する。"""

    try:
        document = Document(path)

        content = ""

        # 通常の段落から本文を取得
        for paragraph in document.paragraphs:
            text = paragraph.text

            if text:
                content += text + "\n"

        # Word内の表からセルの文字を取得
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        content += cell.text + "\n"

        return content

    except Exception as error:
        logging.error(
            "Word読み込み失敗: %s / 理由: %s",
            path,
            error
        )

        return None


def read_excel_file(path):
    """Excelファイルの全シートからセルの値を取得し、検索用本文として返す。"""

    try:
        workbook = load_workbook(path, read_only=True, data_only=True)

        content = ""

        # 全シート・全行・全セルを順番に確認
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        content += str(cell.value) + "\n"

        workbook.close()

        return content

    except Exception as error:
        logging.error(
            "Excel読み込み失敗: %s / 理由: %s",
            path,
            error
        )

        return None


def main():
    """ファイルスキャン、DB更新、検索メニューを実行する。"""

    target_dir = Path(r"C:\Users\deser\Documents")

    logging.info("スキャンを開始します")

    # DBの初期化
    initialize_db()

    # ファイル情報と本文を取得
    files = scan_files(target_dir)

    logging.info("スキャン完了: %d件", len(files))

    # スキャン結果をDBへ保存
    save_files_to_db(files)

    # 実際には存在しなくなったファイルをDBから削除
    delete_missing_files_from_db()

    file_count = count_files_in_db()

    print("DBに保存されているファイル数:", file_count)

    print("1: 拡張子検索")
    print("2: ファイル名検索")
    print("3: ファイル内容検索") 

    choice = input("検索方法を選択してください: ")

    if choice == "1":
        search_extension = input("検索する拡張子を入力してください: ").lower()

        # 「txt」のように入力された場合も「.txt」に統一
        if not search_extension.startswith("."):
            search_extension = "." + search_extension

        results = search_by_extension_db(search_extension)    

    elif choice == "2":
        keyword = input("検索するファイル名を入力してください: ")
        results = search_by_name_db(keyword)

    elif choice == "3":
        keyword = input("検索するファイル内容を入力してください: ")
        results = search_by_content_db(keyword)

    else:
        print("1~3を入力してください。")
        results = []

    if not results:
        print("該当するファイルはありませんでした。")

    else:
        print("検索結果:", len(results), "件")

        for file in results:
            print("ファイル名:", file["name"])
            print("拡張子:", file["extension"])
            print("パス:", file["path"])
            print("サイズ:", file["size"], "bytes")
            print("更新日時:", file["modified"])
            print("-" * 50)        

if __name__ == "__main__":
    main()