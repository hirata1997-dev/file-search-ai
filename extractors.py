import logging
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook


def read_text_file(path):
    """TXTファイルをUTF-8またはCP932で読み込み、本文を返す。"""

    encodings = ["utf-8", "cp932"]

    # UTF-8で失敗した場合はCP932で再試行
    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as file:
                return file.read()

        except UnicodeDecodeError:
            continue
            
    logging.warning(
        "対応している文字コードで読み込めませんでした: %s",
        path
    )

    return None


def read_pdf_file(path):
    """PDFファイルから検索対象となるテキストを抽出する。"""

    try:
        reader = PdfReader(path)

        content = ""

        # PDFをページ単位で読み込み、本文を結合
        for page in reader.pages:
            text = page.extract_text()

            if text:
                content += text + "\n"

        return content

    except Exception as error:
        logging.error(
            "PDF読み込み失敗: %s / 理由: %s",
            path,
            error
        )

        return None


def read_docx_file(path):
    """Wordファイルの段落と表から検索対象となるテキストを抽出する。"""

    try:
        document = Document(path)

        content = ""

        # 通常の段落から本文を取得
        for paragraph in document.paragraphs:
            text = paragraph.text

            if text:
                content += text + "\n"

        # Word内の表からセルの文字を取得
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        content += cell.text + "\n"

        return content

    except Exception as error:
        logging.error(
            "Word読み込み失敗: %s / 理由: %s",
            path,
            error
        )

        return None


def read_excel_file(path):
    """Excelファイルの全シートからセルの値を取得し、検索用本文として返す。"""

    try:
        workbook = load_workbook(path, read_only=True, data_only=True)

        content = ""

        # 全シート・全行・全セルを順番に確認
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        content += str(cell.value) + "\n"

        workbook.close()

        return content

    except Exception as error:
        logging.error(
            "Excel読み込み失敗: %s / 理由: %s",
            path,
            error
        )

        return None